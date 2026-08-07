import io
import re
from pathlib import Path
from typing import Dict

import fitz  # PyMuPDF
from docx import Document

from ingestion.chunker import chunk_documents
from ingestion.embedder import embed_and_store

MAX_FILE_BYTES = 20 * 1024 * 1024  # 20 MB
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract plain text from PDF bytes using PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _extract_text_file(file_bytes: bytes) -> str:
    """Decode plain-text bytes (UTF-8 with a lenient fallback)."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a .docx (paragraphs + table cells) using python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch text extraction by file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf_text(file_bytes)
    if ext == ".docx":
        return _extract_docx_text(file_bytes)
    return _extract_text_file(file_bytes)


def _slugify(name: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")
    return slug or "document"


def ingest_file_bytes(
    filename: str,
    file_bytes: bytes,
    chunk_size: int = 400,
    chunk_overlap: int = 80,
) -> Dict:
    """
    Extract text from an uploaded file (.pdf/.txt/.docx), chunk it, embed it,
    and store it in ChromaDB — mirrors paste_ingest.ingest_paste() but for
    user-uploaded files.
    """
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return {
            "success": False,
            "error": f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}.",
        }

    if len(file_bytes) > MAX_FILE_BYTES:
        return {
            "success": False,
            "error": f"File exceeds the {MAX_FILE_BYTES // (1024 * 1024)}MB upload limit.",
        }

    try:
        text = extract_text(filename, file_bytes)
    except Exception as e:
        return {"success": False, "error": f"Could not read file: {e}"}

    if not text.strip():
        return {
            "success": False,
            "error": (
                f"No extractable text found in '{filename}'. "
                "It may be a scanned/image-only file (OCR isn't supported yet)."
            ),
        }

    stem = Path(filename).stem
    doc_id = _slugify(stem)
    display_name = stem.replace("_", " ").strip() or filename

    document = {
        "doc_id": doc_id,
        "filename": f"{doc_id}{ext}",
        "text": text,
        "source": f"upload:{filename}",
    }

    chunks = chunk_documents([document], chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    embed_and_store(chunks)

    return {
        "success": True,
        "topic": display_name,
        "chunks_added": len(chunks),
    }
